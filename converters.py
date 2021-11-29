"""Contains multiple classes for file conversions. These include:
    * txtconvert
    * imgconvert
    * sheetconvert
"""

from helpers.helpers import typecheck, outpath, compatibility, typencd, sminp, inpchecker
import os
import csv
import cv2
import base64
from PIL import Image
import pandas as pd
import rawpy
from docx import Document
from tqdm import tqdm
from pathlib import Path
from xlsxwriter.workbook import Workbook
import warnings
warnings.filterwarnings("ignore")

class txtconvert:
    """Holds method to convert .txt and other text document files into .docx format. Conversions include

    Args:
        * `__file__` ([type]: str): Input file/directory.
        * `disable` ([type]: bool): Argument to disable or enable the stdout of the tqdm progress bar.
        * `__d__` ([type]: str): Output directory. Default is the current working directory.

    Functions: 

    >>> txt_docx()

    Example with txt_docx():

    >>> from converters import txtconvert

    >>> txtconv = txtconvert(__file__ = path/to/file/, d = path/to/output/directory)

    >>> txtconv.txt_docx()"""

    __d__: str 
    disable: bool
    __file__: str

    def __init__(self, __file__, __disable__ = False, __d__ = os.getcwd()):
        self.__file__ = __file__
        self.disable = __disable__
        self.__d__ = __d__

        # Check class arguments for errors.
        if os.path.isdir(__file__) and len(os.listdir(__file__) ) == 0:
            raise OSError(f'{__file__} directory is empty.')
        elif os.path.isfile(__file__) and os.path.getsize(__file__) == 0:
            raise OSError(f'{__file__} file is empty.')

        if not ((os.path.isdir(__file__)) or os.path.isfile(__file__)):
            raise ValueError('__file__ must be a path to a file or directory.')

        if not isinstance(__disable__, bool):
            raise ValueError(f'__disable__ must be Boolean, not of Type: {(type(__disable__)).__name__}.')

        if not os.path.isdir(__d__):
            raise ValueError(f'__d__ must be a directory and a Type: string, not equal to {__d__} and of Type: {(type(__d__)).__name__}.')

    def txt_docx(self):
        """Convert and Text Document type file or a directory with Text Document type files into `.docx` file/s.
        
        Supported Text Document formats:

            *`.txt`

            *`.log`

            *`.ini`

        Raises:
        * `TypeError`: Raised when file is not a supported text document, or if input is not a file/directory,
        or if input directory does not contain any supported file format.
        """

        inpchecker(inp1 = self.__file__, inp2 = self.__d__, ftype = str)   # Check if objects are strings and 
                                                                            # for the existance of the input paths.

        type_check = typecheck(__object__ = self.__file__)
        extensions = [".txt", ".log", ".ini"] # Supported Text Document file extensions.

        # Check if __file__ instance is a parent/child directory.
        if type_check == True:
            dir_contents = compatibility(__inpobj__ = self.__file__, __compat__ = extensions)
            checktxt = True
            checkfl = None  # __file__ instance is not a single file.
        else:
            type_check = False
            dir_contents = None # __file__ instance is not a parent/child directory.
            checkfl = self.__file__.lower().endswith(extensions) # Check __file__ instance file type.

        # __file__ instance is a parent/child directory and contains at least 1 .txt file.
        if type_check == True and checktxt == True:              
            for f in tqdm(dir_contents, desc = 'Converting %i files to .docx format' %(len(dir_contents)), unit = ' Files', disable = self.disable):  # Iterate over all the entries.
                doc = Document() 
                flpath = os.path.join(self.__file__, f)
                enc = typencd(__inpobj__ = flpath)    
                with open(flpath, 'r', encoding = enc) as inf:    
                    line = inf.read()
                    doc.add_paragraph(line)
                    doc.save(outpath(dinput = self.__d__, flinput = flpath) + ".docx") # full path of output .docx file to save.

        # __file__ instance is *not* a parent/child directory, but is a .txt file.
        elif type_check == False and checkfl == True:
            print(f'Converting {self.__file__} into a .docx format.')
            doc = Document()    
            enc = typencd(__inpobj__ = self.__file__)
            with open(self.__file__, 'r', encoding = enc) as inf:  
                line = inf.read()
                doc.add_paragraph(line)
                output = doc.save(outpath(dinput = self.__d__, flinput = self.__file__) + ".docx") # full path of output .docx file to save.
            print(f'Conversion complete! New file is saved in {output}.')

        # __file__ instance is neither a parent/child directory or a file in .txt format. 
        else:
            raise TypeError(f"{self.__file__} must either be a directory that contains at least 1 compatible file or an individual compatible file.")

class imgconvert():
    """Holds multiple methods to convert image files into other file formats. Conversions include:
       * image to pdf.
       * image to base64 text file.
       * images to binary.
       * multiple compressed/raw image formats to .jpeg and .png

    Args:
        * `__file__` ([type]: str): Input file/directory.
        * `__disable__` ([type]: bool): Argument to disable or enable the stdout of the tqdm progress bar.
        * `__d__` ([type]: str): Output directory. Default is the current working directory.

    Functions:  
                >>> img_pdf()
                
                >>> img_64()
                
                >>> images_binary(bwn = "", keep = True)
                
                >>> img_format(format)

    Example with img_pdf:
                >>> from converters import imgconvert
                
                >>> imgconv = imgconvert(__file__ = path/to/file/or_list_of_file_paths, d = path/to/output/directory)  
                
                >>> imgconv.img_pdf()"""
    
    __d__: str
    __file__: str
    disable: bool
    ext: str
    a: str
    pdf: str
    format: str

    def __init__(self, __file__, __disable__ = False, __d__ = os.getcwd()):
        self.__file__ = __file__
        self.disable = __disable__
        self.__d__ = __d__

        # Check class arguments for errors.
        if os.path.isdir(__file__) and len(os.listdir(__file__) ) == 0:
            raise OSError(f'{__file__} directory is empty.')
        elif os.path.isfile(__file__) and os.path.getsize(__file__) == 0:
            raise OSError(f'{__file__} file is empty.')

        if not ((os.path.isdir(__file__)) or os.path.isfile(__file__)):
            raise ValueError('__file__ must be a path to a file or directory.')

        if not isinstance(__disable__, bool):
            raise ValueError(f'__disable__ must be Boolean, not of Type: {(type(__disable__)).__name__}.')

        if not os.path.isdir(__d__):
            raise ValueError(f'__d__ must be a directory and a Type: string, not equal to {__d__} and of Type: {(type(__d__)).__name__}.')

    def _pdfconv(self, __inp__, __outd__):
        """Inner function to convert a file to pdf.

        Args:
            * `__inp__` ([type]: `str`): Input file
            * `__outd__` ([type]: `str`): Output directory.

        Returns:
            [type]: `str`: Output file path.
        """

        f = Image.open(__inp__)
        Input_File_Name = os.path.splitext(os.path.basename(__inp__))[0]
        subdir = os.path.join(__outd__, Input_File_Name)
        pdf_image = f.convert('RGB')    # Converts RGB values of Image.
        flsave = str(subdir + '.pdf')
        pdf_image.save(flsave)
        return flsave

    def _64conv(self, __inp__, __outd__):
        """Inner function to convert a file to base64 UTF-8 encryption (.txt format).

        Args:
            * `__inp__` ([type]: `str`): Input file.
            * `__outd__` ([type]: `str`): Output directory.

        Returns:
            [type]: `str`: Output file path.
        """

        with open(__inp__, "rb") as img:
            base64_str = base64.b64encode(img.read())
            base64utf8 = base64_str.decode('utf-8') # Encode to UTF-8 standard.
            file_name = os.path.splitext(os.path.basename(__inp__))[0]
            subdir = os.path.join(__outd__, file_name)
            txt_f = (subdir + '.txt')

            with open(txt_f, "w") as txt:
                txt.write(base64utf8)
            return txt_f

    def _imgbnr(self, __inp__, __outd__, __kp__, __bw__):
        """Inner function to convert a file to binary. Original file format type is kept the same.

        Args:
            * `__inp__` ([type]: `str`): Input file.
            * `__outd__` ([type]: `str`): Output directory.
            * `__kp__` ([type]: `bool`): Sets wheter or not to keep the original file. Only used when `__outd__` is the same as the directory of the input file/s.
            * `__bw__` ([type]: `str`): [description]

        Returns:
            [type]: `str`: Path of output file.
        """

        flname, ext = os.path.splitext(__inp__)   # flname is path + flname without extension.
        img = cv2.imread(__inp__, cv2.IMREAD_GRAYSCALE)
        thresh = 128
        img_binary = cv2.threshold(img, thresh, 255, cv2.THRESH_BINARY)[1]

        fname = os.path.basename(__inp__).split('.')[0]     # file name without extension and path. 
        subdir = str(__outd__ + fname)

        if __kp__ == True and subdir == flname and __bw__ == "":
            bwn = "bnr"

        img_bw = str(subdir + bwn + ext)
        cv2.imwrite(img_bw, img_binary)
        return img_bw

    def images_binary(self, bwn = "", keep = True):
        """Convert a single or multiple images to binary. 

        The input image/s must all be located in the same directory. Multiple directories are not supported.

        Args:
            * `bwn` (`str`, `optional`): Rename binary file. Only needed if output directory is the same as input directory. Defaults to an `empty` string.
            * `keep` (`bool`, `optional`): Choose whether to keep the original non-binary file, if output directory is the same as input directory. Defaults to `True`.

        Raises:
        * `TypeError`: Raised when file is not a supported text document, or if input is not a file/directory,
        or if input directory does not contain any supported file format.
        """

        inpchecker(inp1 = self.__file__, inp2 = self.__d__, ftype = str)   # Check if objects are strings and 
                                                                            # for the existance of the input paths.

        type_check = typecheck(__object__ = self.__file__)

        # file formats supported
        extensions = (".jpeg", ".jpg", ".png", ".tif", ".tiff", ".hdr", ".pic", ".sr", ".ras", ".pfm", "pbm", 
                    ".pgm", ".ppm", ".pxm", ".pnm", ".webp", ".jp2", ".jpe", ".bmp", ".dib",
                    ".JPEG",  ".JPG", ".PNG", ".TIF", ".TIFF", ".HDR", ".PIC", ".SR", ".RAS", ".PFM", "PBM", 
                    ".PGM", ".PPM", ".PXM", ".PNM", ".WEBP", ".JP2", ".JPE", ".BMP", ".DIB")

        # Check if __file__ instance is a parent/child directory.
        if type_check == True:
            dir_contents = compatibility(__inpobj__ = self.__file__, __compat__ = extensions)
            checkfl = None  # __file__ instance is not a single file.
        else:
            type_check = False
            dir_contents = None # __file__ instance is not a parent/child directory.
            checkfl = self.__file__.lower().endswith(extensions) # Check __file__ instance file type.

        # __file__ instance is a parent/child directory.
        if type_check == True:
            for f in tqdm(dir_contents, desc = 'Converting %i files to a binary format' %len(dir_contents), unit=' Files', disable = self.disable):  # Iterate over all the entries
                flinp = os.path.join(self.__file__, f)  # Build full path of each iterated input file.
                self._imgbnr(__inp__ = flinp, __outd__ = self.__d__, __kp__ = keep, __bw__ = bwn)

        # __file__ instance is *not* a parent/child directory, but is a supported image file.
        elif type_check == False and checkfl == True:
            print(f'Converting {self.__file__} into a binary format.')
            output = self._imgbnr(__inp__ = self.__file__, __outd__ = self.__d__, __kp__ = keep, __bw__ = bwn)
            print(f'Conversion complete! New file is saved in {output}.')

        # __file__ instance is neither a parent/child directory or a image file. 
        else:
            raise TypeError(f"{self.__file__} must either be a directory that contains at least 1 supported image file or an individual supported image file.")

    def img_base64(self):
        """Encode an image file or directory with images to base64 (`UTF-8`) and save it as `.txt`
        """

        inpchecker(inp1 = self.__file__, inp2 = self.__d__, ftype = str)   # Check if objects are strings and 
                                                                            # for the existance of the input paths.
        type_check = typecheck(__object__ = self.__file__)

        # Supported image extensions for conversion to base64.
        supp_ext = ( '.dng', '.raw', '.cr2', '.crw', '.erf', '.raf', '.tif', '.tiff', '.kdc', '.dcr', '.mos', '.mef', '.nef', '.orf', '.rw2', '.pef', 
                '.x3f', '.srw', '.srf', '.sr2', '.arw', '.mdc', '.bmp', '.mrw', '.DNG', '.RAW', '.CR2', '.CRW', '.ERF', '.RAF', '.TIF', '.TIFF', 
                '.KDC', '.DCR', '.MOS', '.MEF', '.NEF', '.ORF', '.RW2', '.BMP', '.PEF', '.X3F', '.SRW', '.SRF', '.SR2', '.ARW', '.MDC', '.MRW' ,
                '.dng', '.raw', '.cr2', '.crw', '.erf', '.raf', '.kdc', '.dcr', '.mos', '.mef', '.nef', '.orf', '.rw2', '.pef', 
                '.x3f', '.srw', '.srf', '.sr2', '.arw', '.mdc', '.mrw', '.DNG', '.RAW', '.CR2', '.CRW', '.ERF', '.RAF', '.KDC',
                '.DCR', '.MOS', '.MEF', '.NEF', '.ORF', '.RW2', '.PEF', '.X3F', '.SRW', '.SRF', '.SR2', '.ARW', '.MDC', '.MRW', 
                '.jpeg', '.png', '.jpg', '.JPEG', '.PNG', '.JPG')

        # __file__ is a parent/child directory.
        if type_check == True:
            dir_contents = compatibility(__inpobj__ = self.__file__, __compat__ = supp_ext)
            for f in tqdm(dir_contents, desc = 'Converting %i files to base64 format' %len(dir_contents) , unit=' Files', disable = self.disable):  # Iterate over all the entries
                flinp = os.path.join(self.__file__, f)
                self._64conv(__inp__ = flinp, __outd__ = self.__d__) 

        else:
            print(f'Converting {self.__file__} into base64 format with UTF-8 encoding.')
            output = self._64conv(__inp__ = self.__file__, __outd__ = self.__d__)
            print(f'Conversion complete! New file is saved in {output}.')   

    def img_pdf(self):
        """Convert an image file to `.pdf`.

        Supported file format inputs:

        * `.png`, `.jpeg`, `.jpg`, `.dng`, `.raw`, `.cr2`, `.crw`, `.erf`, `.raf`, `.tif`,

        * `.tiff`, `.kdc`, `.dcr`, `.mos`, `.mef`, `.nef`, `.orf`, `.rw2`, `.pef`,

        * `.x3f`, `.srw`, `.srf`, `.sr2`, `.arw`, `.mdc`, `.bmp`, `.mrw`
        """

        inpchecker(inp1 = self.__file__, inp2 = self.__d__, ftype = str)   # Check if objects are strings and 
                                                                            # for the existance of the input paths.

        supp_ext = ( '.dng', '.raw', '.crw', '.erf', '.raf', '.tif', '.tiff', '.kdc', '.dcr', '.mos', '.mef', '.nef', '.orf', '.rw2', '.pef', 
                '.x3f', '.srw', '.srf', '.sr2', '.mdc', '.bmp', '.mrw', '.DNG', '.RAW', '.CRW', '.ERF', '.RAF', '.TIF', '.TIFF', 
                '.KDC', '.DCR', '.MOS', '.MEF', '.NEF', '.ORF', '.RW2', '.BMP', '.PEF', '.X3F', '.SRW', '.SRF', '.SR2', '.MDC', '.MRW' ,
                '.jpeg', '.png', '.jpg', '.JPEG', '.PNG', '.JPG')
        type_check = typecheck(__object__ = self.__file__)

        # __file__ is a parent/child directory.
        if type_check == True:
            dir_contents = compatibility(__inpobj__ = self.__file__, __compat__ = supp_ext)
            for f in tqdm(dir_contents, desc = 'Converting %i files to .pdf format' %len(dir_contents) , unit=' Files', disable = self.disable):  # Iterate over all the entries
                flinp = os.path.join(self.__file__, f)
                self._pdfconv(__inp__ = flinp, __outd__ = self.__d__)

        else:
            print(f'Converting {self.__file__} into .pdf format.')
            output = self._pdfconv(__inp__ = self.__file__, __outd__ = self.__d__)
            print(f'Conversion complete! New file is saved in {output}.')  

    def img_format(self, format):
        """Convert either uncompressed image formats or compressed image formats to `.jpeg` and `.png`.

        Args:
            * `format` ([type]: `str`): format to convert the image to.

        Raises:
            `TypeError`: raises error when an unsupported format is inserted as an input file or
            when an unsupported format is requested in the format input argument.

        Supported file format inputs:

        * .png, .jpeg, .jpg, .dng, .raw, .cr2, .crw, .erf, .raf, .tif,

        * .tiff, .kdc, .dcr, .mos, .mef, .nef, .orf, .rw2, .pef,

        * .x3f, .srw, .srf, .sr2, .arw, .mdc, .bmp, .mrw
        """

        inpchecker(inp1 = self.__file__, inp2 = self.__d__, ftype = str)   # Check if objects are strings and 
                                                                            # for the existance of the input paths.

        ext = ( '.dng', '.raw', '.cr2', '.png', '.jpeg', '.jpg', '.crw', '.erf', '.raf', '.tif', '.tiff', '.kdc', '.dcr', '.mos', '.mef', '.nef', '.orf', '.rw2', '.pef', 
                '.x3f', '.srw', '.srf', '.sr2', '.arw', '.mdc', '.bmp', '.mrw', '.DNG', '.RAW', '.CR2', '.CRW', '.ERF', '.RAF', '.TIF', '.TIFF', 
                '.KDC', '.DCR', '.MOS', '.MEF', '.NEF', '.ORF', '.RW2', '.BMP', '.PEF', '.X3F', '.SRW', '.SRF', '.SR2', '.ARW', '.MDC', '.MRW', '.PNG', '.JPEG', '.JPG')

        # Doesn't contain .tiff, .tif, .bmp extensions.
        raw_ext = ( '.dng', '.raw', '.cr2', '.crw', '.erf', '.raf', '.kdc', '.dcr', '.mos', '.mef', '.nef', '.orf', '.rw2', '.pef', 
                '.x3f', '.srw', '.srf', '.sr2', '.arw', '.mdc', '.mrw', '.DNG', '.RAW', '.CR2', '.CRW', '.ERF', '.RAF', '.KDC',
                '.DCR', '.MOS', '.MEF', '.NEF', '.ORF', '.RW2', '.PEF', '.X3F', '.SRW', '.SRF', '.SR2', '.ARW', '.MDC', '.MRW' )

        supp_ext = ('.jpeg','.png', '.jpg')

        if not format in supp_ext:
            raise TypeError(f"{format} file format not supported.")

        def _tojpg(flname, infl):   # Used in type_check if and elif statements.
            """Local function that converts the input image file to .jpeg/.jpg.
            
            Returns:
                ([type]: `str`): The full path of the output .jpeg/.jpg file."""

            if infl.endswith(raw_ext):
                outfl = flname + ".jpeg"
                with rawpy.imread(infl) as raw:
                    rgb = raw.postprocess(use_auto_wb = True)
                Image.fromarray(rgb).save(outfl, quality = 90, optimize = True)
                raw.close()
            else:
                outfl = flname + ".jpeg"
                im = Image.open(infl)
                out = im.convert("RGB")
                out.save(outfl, "jpeg", quality=100, subsampling=2)
            return outfl

        def _topng(flname, infl):   # Used in type_check if and elif statements.
            """Local function that converts the input image file to .png

            Returns:
                ([type]: `str`): The full path of the output .png file."""

            if infl.endswith(raw_ext):
                outfl = flname + ".png"
                with rawpy.imread(infl) as raw:
                    rgb = raw.postprocess(use_auto_wb = True)
                Image.fromarray(rgb).save(outfl, quality = 90, optimize = True)
                raw.close()
            else:
                outfl = flname + ".png"
                im = Image.open(infl)
                out = im.convert("RGB")
                out.save(outfl, "png", quality=100, subsampling=2)
            return outfl

        def _loop_flobj(obj, fl, adir): # Used in for loop in directory contents list comprehension.
            """Local function that assigns the name and path of the output file.

            Returns:
                ([type]: `str`): The full path of the output file [0] and the full path of the input file [1]."""

            flpath = os.path.join(obj, fl)
            file_name = os.path.splitext(os.path.basename(flpath))[0]
            subdir = os.path.join(adir, file_name)

            return subdir, flpath

        type_check = typecheck(__object__ = self.__file__)

        # __file__ is a parent/child directory.
        if type_check == True:
            dir_contents = compatibility(__inpobj__ = self.__file__, __compat__ = ext)
            for f in tqdm(dir_contents, desc = 'Converting %i files to %s format' %(len(dir_contents),format) , unit=' Files', disable = self.disable):  # Iterate over all the entries
                fl_handler = _loop_flobj(obj = self.__file__, fl = f, adir = self.__d__)
                var, obj_ext = os.path.splitext(f)    # var is placeholder, obj_ext is the input file extension.
                if f.endswith(ext) and format == ".jpeg" and obj_ext !=".jpeg" and obj_ext !=".jpg":   # Desired format is JPEG.
                    _tojpg(flname = fl_handler[0], infl = fl_handler[1])

                elif f.endswith(ext) and format == ".jpg" and obj_ext !=".jpeg" and obj_ext !=".jpg":   # Desired format is JPG.
                    _tojpg(flname = fl_handler[0], infl = fl_handler[1])

                elif f.endswith(ext) and format == ".png" and obj_ext !=".png":    # Desired format is PNG.
                    _topng(flname = fl_handler[0], infl = fl_handler[1])

        # __file__ is not a parent/child directory, desired format is JPEG.
        elif type_check == False and self.__file__.endswith(ext) and format == ".jpeg":
            subdir = outpath(dinput = self.__d__, flinput = self.__file__)
            print(f'Converting {self.__file__} into .jpeg format.')        
            output = _tojpg(flname = subdir, infl = self.__file__)
            print(f'Conversion complete! New file is saved in {output}.')  
        
        # __file__ is not a parent/child directory, desired format is JPG.
        elif type_check == False and self.__file__.endswith(ext) and format == ".jpg":        
            subdir = outpath(dinput = self.__d__, flinput = self.__file__)
            print(f'Converting {self.__file__} into .jpg format.') 
            output = _tojpg(flname = subdir, infl = self.__file__)
            print(f'Conversion complete! New file is saved in {output}.')  

        # __file__ is not a parent/child directory, desired format is PNG.
        elif type_check == False and self.__file__.endswith(ext) and format == ".png":        
            subdir = outpath(dinput = self.__d__, flinput = self.__file__)
            print(f'Converting {self.__file__} into .png format.') 
            output = _topng(flname = subdir, infl = self.__file__)
            print(f'Conversion complete! New file is saved in {output}.')

        # __file__ format is not supported.
        else:
            var, obj_ext = os.path.splitext(self.__file__)    # var is placeholder, obj_ext is the input file extension.
            raise TypeError(f'{obj_ext} file format of {self.__file__} input file is not supported.')

class sheetconvert():
    """Allows for the conversion of a single or multiple spreadsheet files
    (.xlsx, .csv, .tsv) into .xlsx, .csv or .tsv.

    Args:
        * `__file__` ([type]: str): Input file/directory.
        * `disable` ([type]: bool): Argument to disable or enable the stdout of the tqdm progress bar.
        * `__d__` ([type]: str): Output directory. Default is the current working directory.

    Functions:
        >>> convertsh(totype)
    
    Example:
        >>> from converters import sheetconvert
        >>> a = sheetconvert(__file__ = 'path/to/file/or/dir', __d__ = 'path/to/output/dir')
        >>> a.convertsh(totype = '.csv')
    """

    __file__: str
    __d__: str
    disable: bool
    spreadsheet: str
    csv_file: str
    xlsx_name: str 

    def __init__(self, __file__, __disable__ = False, __d__ = os.getcwd()):
        self.__file__ = __file__
        self.disable = __disable__
        self.__d__ = __d__

        # Check class arguments for errors.
        if os.path.isdir(__file__) and len(os.listdir(__file__) ) == 0:
            raise OSError(f'{__file__} directory is empty.')
        elif os.path.isfile(__file__) and os.path.getsize(__file__) == 0:
            raise OSError(f'{__file__} file is empty.')

        if not ((os.path.isdir(__file__)) or os.path.isfile(__file__)):
            raise ValueError('__file__ must be a path to a file or directory.')

        if not isinstance(__disable__, bool):
            raise ValueError(f'__disable__ must be Boolean, not of Type: {(type(__disable__)).__name__}.')

        if not os.path.isdir(__d__):
            raise ValueError(f'__d__ must be a directory and a Type: string, not equal to {__d__} and of Type: {(type(__d__)).__name__}.')

    def _xlsx_csv(self, spreadsheet, direc):
            """Internal function to convert an `.xslx` file to a `.csv` file.

            Args:
                * `spreadsheet` ([type]: `str`): path to the .xlsx file.
                * `direc` ([type]: `str`): Output directory of the .csv file.

            Returns:
                ([type]: `str`): The path of the output .csv.
            """

            xlsx = pd.read_excel(spreadsheet)
            file_name = os.path.splitext(os.path.basename(spreadsheet))[0]
            subdir = os.path.join(direc, file_name)
            csv_file = str(subdir + '.csv')
            xlsx.to_csv (csv_file, 
                    index = None,
                    header=True)

            return csv_file

    def _csv_xlsx(self, spreadsheet, direc):
        """Internal function to convert a `.csv` file to an `.xlsx` file.

        Args:
            * `spreadsheet` ([type]: `str`): path to the .csv file.
            * `direc` ([type]: `str`): Output directory of the .xlsx file.

        Returns:
            ([type]: `str`): The path of the output .xlsx.
        """

        with open(spreadsheet, "r"):
            read_file = pd.read_csv(spreadsheet)
            file_name = os.path.splitext(os.path.basename(spreadsheet))[0]
            subdir = os.path.join(direc, file_name)

            xlsx_name = str(subdir + ".xlsx")
            xlsx = pd.ExcelWriter(xlsx_name)
            read_file.to_excel(xlsx, index = False)
            xlsx.save()

            return xlsx_name

    def _xlsx_tsv(self, spreadsheet, direc):
        """Internal function to convert an `.xlsx` file to a `.tsv` file.

        Args:
            * `spreadsheet` ([type]: `str`): path to the .xlsx file.
            * `direc` ([type]: `str`): Output directory of the .tsv file.

        Returns:
            ([type]: `str`): The path of the output .tsv.
        """

        fname = Path(spreadsheet).stem
        outname = os.path.join(direc, fname)
        outtsv = str(outname + ".tsv")
        data_xlsx = pd.read_excel(spreadsheet, index_col=None)

        df = data_xlsx.replace('\n', ' ',regex=True)
        df.to_csv(outtsv, sep='\t', encoding='utf-8',  index=False, line_terminator='\r\n')

        return outtsv

    def _tsv_xlsx(self, spreadsheet, direc):
        """Internal function to convert a `.tsv` file to an `.xlsx` file.

        Args:
            * `spreadsheet` ([type]: `str`): path to the .tsv file.
            * `direc` ([type]: `str`): Output directory of the .xlsx file.

        Returns:
            ([type]: `str`): The path of the output .xlsx.
        """

        # Create workbook with the appropriate file path.
        xlsxname = os.path.splitext(os.path.basename(spreadsheet))[0]
        xlsxnmext = xlsxname + ".xlsx" 
        xlsxp = os.path.join(direc, xlsxnmext)
        workbook = Workbook(xlsxp)
        worksheet = workbook.add_worksheet()

        read_tsv = csv.reader(open(spreadsheet, 'r', encoding='utf-8'), delimiter='\t')
        for row, data in enumerate(read_tsv):   # loop through the row data in the tsv.
            worksheet.write_row(row, 0, data)
        workbook.close()

        return xlsxp

    def _csv_tsv_csv(self, spreadsheet, direc):
        """Internal function to convert a `.csv` file to a `.tsv` file and vice versa.

        Args:
            * `spreadsheet` ([type]: `str`): path to the .csv/.tsv file.
            * `direc` ([type]: `str`): Output directory of the .tsv/.csv file.

        Returns:
            ([type]: `str`): The path of the output .tsv/.csv.
        """

        # Check if the input file is a .tsv or a .csv file.
        if spreadsheet.endswith('.tsv') or spreadsheet.endswith('.TSV'):
            csvname = os.path.splitext(os.path.basename(spreadsheet))[0]
            csvnmext = csvname + ".tsv" 
            csvp = os.path.join(direc, csvnmext)
            
            with open(spreadsheet,'r', encoding='utf-8') as tsvin, open(csvp, 'w', newline='', encoding='utf-8') as csvout:
                tsvin = csv.reader(tsvin)
                csvout = csv.writer(csvout, delimiter='\t')
                for row in tsvin:
                    csvout.writerow(row)
            return csvp

        elif spreadsheet.endswith('.csv') or spreadsheet.endswith('.CSV'):
            tsvname = os.path.splitext(os.path.basename(spreadsheet))[0]
            tsvnmext = tsvname + ".csv" 
            tsvp = os.path.join(direc, tsvnmext)
            
            with open(spreadsheet,'r', encoding='utf-8') as csvin, open(tsvp, 'w', newline='', encoding='utf-8') as tsvout:
                csvin = csv.reader(csvin)
                tsvout = csv.writer(tsvout, delimiter='\t')
                for row in csvin:
                    tsvout.writerow(row)
            return tsvp

    def _conversion_method(self, __inp__, outdir, typeinp):
        """Internal function. Checks for which conversion type the user requests and correlates that with the input file type.

        Args:
            * `__inp__` ([type]: `str`): Input file.
            * `outdir` ([type]): Output directory.
            * `typeinp` ([type]: `str`): Type of output file.
        """

        # Supported extensions.
        xlsxf = [".xlsx", ".XLSX"]
        csvf = [".csv", ".CSV"]
        tsvf = [".tsv", ".TSV"]

        # Correlate input file type with requested file conversion type. Run the appropriate conversion function.
        if __inp__.endswith(tuple(xlsxf)) and typeinp == ".csv":
            output = self._xlsx_csv(spreadsheet = __inp__, direc = outdir)

        elif __inp__.endswith(tuple(xlsxf)) and typeinp == ".tsv":
            output = self._xlsx_tsv(spreadsheet = __inp__, direc = outdir)

        elif __inp__.endswith(tuple(csvf)) and typeinp == ".xlsx":
            output = self._csv_xlsx(spreadsheet = __inp__, direc = outdir)

        elif __inp__.endswith(tuple(csvf)) and typeinp == ".tsv":
            output = self._csv_tsv_csv(spreadsheet = __inp__, direc = outdir)

        elif __inp__.endswith(tuple(tsvf)) and typeinp == ".xlsx":
            output = self._tsv_xlsx(spreadsheet = __inp__, direc = outdir)

        elif __inp__.endswith(tuple(tsvf)) and typeinp == ".csv":
            output = self._csv_tsv_csv(spreadsheet = __inp__, direc = outdir)

        return output

    def convertsh(self, fromtype, totype):
        """Convert an `.xlsx`/`.csv`/`.tsv` file to either `.xlsx` or `.csv` or `.tsv`. If the file is
        directory, it converts all the files in the directory that are `.xlsx`, `.tsv` and `.csv`.

        Args:
            * `fromtype` ([type]:`str`): specify what type of file to convert, e.g: '.csv', '.xlsx', '.tsv'. Used `only` on batch conversions.
            * `totype` ([type]:`str`): type of converion, e.g: '.csv', '.xlsx', '.tsv'.
        """

        inpchecker(inp1 = self.__file__, inp2 = self.__d__, ftype = str)   # Check if objects are strings and 
                                                                            # for the existance of the input paths.

        supp_ext = ( ".xlsx", ".csv", ".tsv" )  # Spreadsheet supported extenstions.

        if not fromtype in supp_ext:
            raise TypeError(f"{fromtype} file extension not supported")

        if not totype in supp_ext:
            raise TypeError(f"{totype} file extension not supported")

        type_check = typecheck(__object__ = self.__file__)

        # __file__ is a parent/child directory.
        if type_check == True:
            dir_contents = compatibility(__inpobj__ = self.__file__, __compat__ = supp_ext)
            dir_contents[:] = [fl for fl in dir_contents if any(ext not in fl for ext in totype)]   # Removes elements in directory that contain the selected extension.
            dir_contents[:] = [fl for fl in dir_contents if any(ext in fl for ext in fromtype)]   # Removes elements in directory that don't contain the selected extension.
            for f in tqdm(dir_contents, desc = 'Converting %i files to %s format' %(len(dir_contents),totype) , unit=' Files', disable = self.disable):  # Iterate over all the entries
                build_fpath = os.path.join(self.__file__, f)
                self._conversion_method(__inp__ = build_fpath, outdir = self.__d__, typeinp = totype)

        # __file__ is not a parent/child directory.
        else:
            sminp(self.__file__, totype)   # Check if input file extension and 
                                            # totype extension are the same.
            print(f'Converting {self.__file__} into {totype} format.') 
            output = self._conversion_method(__inp__ = self.__file__, outdir = self.__d__, typeinp = totype)
            print(f'Conversion complete! New file is saved in {output}.')